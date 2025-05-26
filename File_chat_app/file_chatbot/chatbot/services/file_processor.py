import os
from typing import List, Dict, Any, Tuple
import re
import pandas as pd
from PyPDF2 import PdfReader
import docx
import logging

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("document_extraction.log"),
        logging.StreamHandler()
    ]
)

logger = logging.getLogger("document_extraction")

def extract_text_from_file(file_path: str) -> str:
    """Extract text from various file types"""
    if not os.path.exists(file_path):
        logger.error(f"File not found: {file_path}")
        raise FileNotFoundError(f"File not found: {file_path}")
    
    file_extension = os.path.splitext(file_path)[1].lower()
    file_size = os.path.getsize(file_path) / (1024 * 1024)  # Size in MB
    
    logger.info(f"Starting text extraction for file: {file_path}")
    logger.info(f"File type: {file_extension}, Size: {file_size:.2f} MB")
    
    try:
        if file_extension == '.pdf':
            text = extract_text_from_pdf(file_path)
        elif file_extension == '.txt':
            text = extract_text_from_txt(file_path)
        elif file_extension == '.docx':
            text = extract_text_from_docx(file_path)
        elif file_extension in ['.csv', '.xlsx', '.xls']:
            text = extract_text_from_tabular(file_path)
        else:
            logger.error(f"Unsupported file type: {file_extension}")
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        logger.info(f"Successfully extracted text from {file_path}")
        logger.debug(f"Extracted text length: {len(text)} characters")
        return text
    
    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {str(e)}", exc_info=True)
        raise

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file"""
    logger.info(f"Extracting text from PDF: {file_path}")
    text = ""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PdfReader(file)
            total_pages = len(pdf_reader.pages)
            logger.info(f"PDF has {total_pages} pages")
            
            for i, page in enumerate(pdf_reader.pages):
                logger.debug(f"Processing page {i+1}/{total_pages}")
                page_text = page.extract_text()
                text += page_text + "\n"
                
                # Log warning if page has little or no text
                if len(page_text.strip()) < 10:
                    logger.warning(f"Page {i+1} appears to have little or no extractable text")
        
        logger.info(f"PDF extraction complete for {file_path}")
        return text
    except Exception as e:
        logger.error(f"Failed to extract text from PDF {file_path}: {str(e)}")
        raise

def extract_text_from_txt(file_path: str) -> str:
    """Extract text from TXT file"""
    logger.info(f"Extracting text from TXT file: {file_path}")
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        logger.info(f"TXT extraction complete for {file_path}")
        return text
    except UnicodeDecodeError:
        logger.warning(f"UTF-8 decoding failed for {file_path}, trying with different encodings")
        try:
            # Try with a different encoding if UTF-8 fails
            with open(file_path, 'r', encoding='latin-1') as file:
                text = file.read()
            logger.info(f"TXT extraction complete using latin-1 encoding for {file_path}")
            return text
        except Exception as e:
            logger.error(f"Failed to extract text from TXT file {file_path}: {str(e)}")
            raise
    except Exception as e:
        logger.error(f"Failed to extract text from TXT file {file_path}: {str(e)}")
        raise

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file"""
    logger.info(f"Extracting text from DOCX file: {file_path}")
    try:
        doc = docx.Document(file_path)
        paragraph_count = len(doc.paragraphs)
        logger.info(f"DOCX has {paragraph_count} paragraphs")
        
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        # Log warning if document has little text
        if len(text.strip()) < 50:
            logger.warning(f"DOCX file {file_path} appears to have little content")
            
        logger.info(f"DOCX extraction complete for {file_path}")
        return text
    except Exception as e:
        logger.error(f"Failed to extract text from DOCX {file_path}: {str(e)}")
        raise

def extract_text_from_tabular(file_path: str) -> str:
    """Extract text from CSV or Excel file"""
    file_extension = os.path.splitext(file_path)[1].lower()
    logger.info(f"Extracting text from tabular file ({file_extension}): {file_path}")
    
    try:
        if file_extension == '.csv':
            logger.debug("Reading CSV file")
            try:
                df = pd.read_csv(file_path)
            except UnicodeDecodeError:
                logger.warning("UTF-8 decoding failed, trying with different encoding")
                df = pd.read_csv(file_path, encoding='latin-1')
        else:  # Excel files
            logger.debug("Reading Excel file")
            df = pd.read_excel(file_path)
        
        rows, cols = df.shape
        logger.info(f"Tabular file has {rows} rows and {cols} columns")
        
        # Check for empty dataframe
        if rows == 0 or cols == 0:
            logger.warning(f"Tabular file {file_path} appears to be empty")
        
        # Convert DataFrame to a text representation
        text = df.to_string(index=False)
        logger.info(f"Tabular extraction complete for {file_path}")
        return text
    except Exception as e:
        logger.error(f"Failed to extract text from tabular file {file_path}: {str(e)}")
        raise